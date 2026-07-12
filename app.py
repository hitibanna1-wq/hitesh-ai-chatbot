import io
import re
import base64
import streamlit as st
from datetime import date
from ddgs import DDGS
from groq import Groq
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from xml.sax.saxutils import escape
from pypdf import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

st.set_page_config(
    page_title="Nivora AI",
    page_icon="✦",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
html, body, [class*="css"] {font-family:'Inter',sans-serif;}
.stApp {background:linear-gradient(145deg,#f8faff 0%,#eef3ff 55%,#f7f9fc 100%);color:#15213a;}
.stApp p,.stApp li,.stApp label,.stApp h1,.stApp h2,.stApp h3,.stApp span {color:#15213a;}
[data-testid="stSidebar"] {background:#ffffff;border-right:1px solid #dce3f0;}
[data-testid="stHeader"] {background:transparent;}
.block-container {max-width:960px;padding-top:2rem;padding-bottom:7rem;}
.brand {font-size:1.35rem;font-weight:700;letter-spacing:-.03em;margin-bottom:.2rem;}
.brand-dot {color:#5865e8}.side-note {color:#65718a;font-size:.78rem;margin-bottom:1.5rem;}
.hero {text-align:center;padding:3rem 1rem 2rem;}
.hero-badge {display:inline-block;color:#4a54c7;background:#e8ebff;border:1px solid #cbd1ff;
padding:.35rem .75rem;border-radius:999px;font-size:.76rem;font-weight:600;letter-spacing:.05em;}
.hero h1 {font-size:clamp(2.2rem,6vw,4rem);line-height:1.05;letter-spacing:-.055em;margin:1rem 0 .7rem;
background:linear-gradient(100deg,#111a31 20%,#5462d9 60%,#128c84);-webkit-background-clip:text;color:transparent;}
.hero p {color:#65718a;font-size:1.05rem;margin:auto;max-width:620px;}
.suggestion {padding:1rem 1.1rem;border:1px solid #d8dfed;border-radius:14px;background:#ffffff;
color:#17213a;min-height:88px;font-size:.9rem;box-shadow:0 8px 28px rgba(54,70,110,.08);}
[data-testid="stChatMessage"] {background:#ffffff;border:1px solid #dce3ef;border-radius:18px;padding:1rem;box-shadow:0 5px 18px rgba(54,70,110,.06);}
[data-testid="stChatInput"] {border:1px solid #b9c4d9;border-radius:18px;background:#ffffff;}
.stButton>button,.stDownloadButton>button {border-radius:12px;border:1px solid #c7d0e2;background:#ffffff;color:#25304a;}
.stButton>button:hover {border-color:#7c8cff;color:white;}
hr {border-color:#e0e6f0}.status {color:#168678;font-size:.78rem;}
</style>
""", unsafe_allow_html=True)

client = Groq(api_key=st.secrets["GROQ_API_KEY"])

SYSTEM_PROMPT = """You are Nivora, a polished professional AI productivity assistant.
Answer in the same language as the user; understand Hindi, Hinglish, and English naturally.
Lead with the direct answer. Be accurate, practical, calm, and concise unless detail is requested.
Use short headings and bullets only when they improve clarity. Never use filler or exaggerated claims.
For instructions, give ordered, actionable steps. For business writing, produce ready-to-use copy.
When uncertain, say what is uncertain. Never invent facts, links, citations, or live information.
When web results are provided, rely on them, cite them as [1], [2], and separate facts from inference.
For risky medical, legal, or financial topics, state appropriate limitations and recommend a professional.
"""

def make_pdf(text: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=20*mm, leftMargin=20*mm,
                            topMargin=18*mm, bottomMargin=18*mm,
                            title="Nivora AI Response", author="Nivora AI")
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("NivoraTitle", parent=styles["Title"], fontName="Helvetica-Bold",
                                 fontSize=20, textColor=HexColor("#25304A"), alignment=TA_CENTER, spaceAfter=14)
    body_style = ParagraphStyle("NivoraBody", parent=styles["BodyText"], fontName="Helvetica",
                                fontSize=10.5, leading=16, textColor=HexColor("#263248"), spaceAfter=8)
    story = [Paragraph("Nivora AI", title_style), Spacer(1, 4*mm)]
    clean = re.sub(r"[*#`]", "", text)
    for block in clean.split("\n"):
        if block.strip():
            safe = escape(block.strip()).encode("latin-1", "replace").decode("latin-1")
            story.append(Paragraph(safe, body_style))
    doc.build(story)
    return buffer.getvalue()

def make_docx(text: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Nivora AI")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(37, 48, 74)
    subtitle = document.add_paragraph("Professional AI Document")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.color.rgb = RGBColor(101, 113, 138)
    document.add_paragraph()
    clean = re.sub(r"[*`]", "", text)
    for line in clean.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("#"):
            document.add_heading(line.lstrip("# "), level=min(line.count("#"), 3))
        elif line.startswith(("- ", "• ")):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.line_spacing = 1.15
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()

def requested_file_type(text: str):
    lowered = text.lower()
    if any(word in lowered for word in ["pdf banao", "pdf bana", "pdf do", "as pdf", "pdf file"]):
        return "pdf"
    if any(word in lowered for word in ["word banao", "word file", "docx", "word document"]):
        return "docx"
    return None

def extract_file_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)[:30000]
    if name.endswith(".docx"):
        document = Document(io.BytesIO(data))
        return "\n".join(p.text for p in document.paragraphs)[:30000]
    return data.decode("utf-8", errors="replace")[:30000]

def transcribe_audio(audio_file) -> str:
    result = client.audio.transcriptions.create(
        file=("voice.wav", audio_file.getvalue()),
        model="whisper-large-v3-turbo",
        response_format="json",
    )
    return result.text

if "messages" not in st.session_state:
    st.session_state.messages = []

with st.sidebar:
    st.markdown('<div class="brand"><span class="brand-dot">✦</span> Nivora AI</div><div class="side-note">Your intelligent workspace</div>', unsafe_allow_html=True)
    if st.button("＋  New conversation", use_container_width=True):
        st.session_state.messages = []
        st.rerun()
    st.divider()
    web_search = st.toggle("Web intelligence", help="Search the live web before answering")
    response_style = st.selectbox("Response style", ["Balanced", "Concise", "Detailed"])
    st.divider()
    st.markdown('<div class="status">● Online & ready</div>', unsafe_allow_html=True)
    st.caption("AI can make mistakes. Verify important information.")

if not st.session_state.messages:
    st.markdown("""
    <div class="hero">
      <span class="hero-badge">YOUR AI WORKSPACE</span>
      <h1>What will you create today?</h1>
      <p>Research, write, analyse and turn ideas into polished work—through one intelligent conversation.</p>
    </div>
    """, unsafe_allow_html=True)
    cols = st.columns(3)
    cards = [
        ("Research", "Compare the best tools for a small online business"),
        ("Create", "Write a professional proposal for my client"),
        ("Plan", "Build a practical 30-day launch strategy"),
    ]
    for col, (title, text) in zip(cols, cards):
        col.markdown(f'<div class="suggestion"><b>{title}</b><br><span style="color:#8792a8">{text}</span></div>', unsafe_allow_html=True)

with st.expander("Add files, images or voice", expanded=False):
    uploaded_files = st.file_uploader(
        "Attach files",
        type=["png", "jpg", "jpeg", "webp", "pdf", "docx", "txt", "csv"],
        accept_multiple_files=True,
        help="Upload up to 5 images or supporting documents.",
    )
    voice_recording = st.audio_input("Record a voice message")
    send_voice = st.button("Transcribe and send voice", disabled=voice_recording is None)

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

prompt = st.chat_input("Ask Nivora anything...")

if send_voice and voice_recording:
    try:
        with st.spinner("Understanding your voice..."):
            prompt = transcribe_audio(voice_recording)
    except Exception as error:
        st.error(f"Voice transcription failed: {error}")
        prompt = None

if prompt:
    export_type = requested_file_type(prompt)
    previous_answer = next(
        (item["content"] for item in reversed(st.session_state.messages) if item["role"] == "assistant"),
        None,
    )
    refers_to_previous = any(word in prompt.lower() for word in ["iska", "iss answer", "this answer", "upar wale", "previous"])
    direct_export_text = previous_answer if export_type and refers_to_previous and previous_answer else None
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    search_results = []
    model_messages = [{"role": "system", "content": SYSTEM_PROMPT}]
    model_messages.extend(st.session_state.messages)

    image_parts = []
    document_context = []
    for uploaded in uploaded_files or []:
        if uploaded.type and uploaded.type.startswith("image/"):
            encoded = base64.b64encode(uploaded.getvalue()).decode("ascii")
            image_parts.append({
                "type": "image_url",
                "image_url": {"url": f"data:{uploaded.type};base64,{encoded}"},
            })
        else:
            document_context.append(f"FILE: {uploaded.name}\n{extract_file_text(uploaded)}")

    if document_context:
        model_messages[-1] = {
            "role": "user",
            "content": prompt + "\n\nAttached document content:\n" + "\n\n".join(document_context),
        }

    if image_parts:
        text_for_image = prompt
        if document_context:
            text_for_image += "\n\nAttached document content:\n" + "\n\n".join(document_context)
        model_messages[-1] = {
            "role": "user",
            "content": [{"type": "text", "text": text_for_image}, *image_parts[:5]],
        }

    with st.chat_message("assistant"):
        try:
            if direct_export_text:
                answer = "Your requested file is ready."
            elif web_search:
                with st.spinner("Searching reliable sources..."):
                    search_results = list(DDGS().text(prompt, region="in-en", safesearch="moderate", max_results=6))
                web_context = "\n\n".join(
                    f"[{i}] {item.get('title','')}\n{item.get('body','')}\nURL: {item.get('href','')}"
                    for i, item in enumerate(search_results, 1)
                )
                model_messages[-1] = {
                    "role": "user",
                    "content": f"Date: {date.today().isoformat()}\n\nQuestion: {prompt}\n\nWeb results:\n{web_context}\n\nAnswer using these sources and cite factual claims [1], [2].",
                }

            if not direct_export_text:
                style_instruction = {
                    "Concise": "Keep the answer brief and focused.",
                    "Detailed": "Give a thorough, structured answer with useful context.",
                    "Balanced": "Use a balanced level of detail.",
                }[response_style]
                model_messages.append({"role": "system", "content": style_instruction})

                with st.spinner("Thinking..."):
                    response = client.chat.completions.create(
                        model="meta-llama/llama-4-scout-17b-16e-instruct" if image_parts else "llama-3.3-70b-versatile",
                        messages=model_messages,
                        temperature=0.35,
                    )
                answer = response.choices[0].message.content
            st.markdown(answer)
            export_content = direct_export_text or answer
            if export_type == "pdf":
                st.download_button("Download PDF", make_pdf(export_content), "nivora-document.pdf", "application/pdf", use_container_width=True)
            elif export_type == "docx":
                st.download_button("Download Word document", make_docx(export_content), "nivora-document.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

            if search_results:
                with st.expander("Sources"):
                    for i, item in enumerate(search_results, 1):
                        st.markdown(f"{i}. [{item.get('title','Source')}]({item.get('href','')})")
        except Exception as error:
            answer = "I couldn't complete that request right now. Please try again in a moment."
            st.error(answer)
            st.caption(f"Technical detail: {error}")

    st.session_state.messages.append({"role": "assistant", "content": answer})
